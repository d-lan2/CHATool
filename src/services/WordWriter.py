from os import write
from docx import Document
from docx.shared import Inches
from ..classes import Output

class WordWriter:
    document = None

    def writeDoc(self,filePath, headerResults : Output.HeaderResult = None, cookieResults: Output.CookieResults = None):
        self.cloneTemplate(filePath)
        if headerResults:
            self.writeHeaders(headerResults)
        if cookieResults:
            self.writeCookies(cookieResults)
        self.document.save(filePath)
    
    def writeHeaders(self,results: Output.HeaderResult):
        headerTable = self.document.tables[1]
        for key in results.headersReportData:
            newTableRowCells = headerTable.add_row().cells
            newTableRowCells[0].text = key
            newTableRowCells[1].text = results.headersReportData[key]

    def writeCookies(self,results: Output.CookieResults):
        cookieTable = self.document.tables[3]
        for key in results.recommendations:
            newTableRowCells = cookieTable.add_row().cells
            newTableRowCells[0].text = key
            newTableRowCells[1].text = results.recommendations[key]

    def cloneTemplate(self,filePath):
        try:
            document = Document('src\\assets\\HeadersTemplate.docx')
        except:
             print("An exception occurred opening HeadersTemplate.docx") 

        try:
            document.save(filePath)
        except:
             print("Cannot save doc to specified locaiton")

        self.document = document