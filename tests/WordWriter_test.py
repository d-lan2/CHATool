from ..src.services.WordWriter import WordWriter
from ..src.services.Auditor import HeaderAuditor
from ..src.classes.Output import HeaderResult
from docx import Document
from pytest_mock import mocker
import os

#Tests written in the AAA format. See HTTP_test.py for more info
def test_can_open_and_save_new_copy_of_template(mocker) -> None:
    #Arrange
    testFilePath = 'output\\test1.docx'
    fakeResult = HeaderResult()
  
    #Act
    WordWriter().writeDoc(testFilePath,fakeResult)

    #Assert
    #Asseting the conents of two word documents in a unit test seems like a real pain. Though it can be done my manually inspecting the xml
    #Perhaps an alternative is to just compare the hashes of the files. See: https://www.pythoncentral.io/hashing-files-with-python/
    #Comparing the files with filecomp didnt work
    #For now just check that the file exists after the test is ran
    assert os.path.isfile(testFilePath) == True

    #temp teardown - TODO impliment proper setup/teardown functions
    os.remove(testFilePath)  
    

def test_can_add_headers_findings_to_report(mocker) -> None:
    #Arrange
    fakeResults =  mocker.MagicMock()
    fakeResults.headersReportData = {"Content-Security-Policy": "Some description", "X-Content-Type-Options":"Some description 2"}
    testFilePath = 'output\\test2.docx'

    #Act
    WordWriter().writeDoc(testFilePath,fakeResults)

    #Assert
    if os.path.isfile(testFilePath):
        document = Document(testFilePath)
    assert os.path.isfile(testFilePath) == True
    assert document.tables[1].rows[1].cells[0].text == "Content-Security-Policy"
    assert document.tables[1].rows[2].cells[1].text == "Some description 2"

    #temp teardown - TODO impliment proper setup/teardown functions
    os.remove(testFilePath)  

def test_can_add_cookie_findings_to_report(mocker) -> None:
    #Arrange
    fakeCookieResults =  mocker.MagicMock()
    fakeCookieResults.recommendations = {"Some cookie attribute": "Some description"}
    testFilePath = 'output\\test3.docx'

    #Act
    WordWriter().writeDoc(testFilePath,None,fakeCookieResults)

    #Assert
    if os.path.isfile(testFilePath):
        document = Document(testFilePath)
    assert os.path.isfile(testFilePath) == True
    assert document.tables[3].rows[1].cells[0].text == "Some cookie attribute"
    assert document.tables[3].rows[1].cells[1].text == "Some description"

    #temp teardown - TODO impliment proper setup/teardown functions
    os.remove(testFilePath)  