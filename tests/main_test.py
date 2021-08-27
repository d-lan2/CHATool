from ..main import main, parameterisedMain
from docx import Document
import os

#Poor and unreliable integration test im using just for devving
#if the site changes its headers then this may fail
def test_can_access_main(mocker) -> None:
    #Arrange
    testFilePath = 'output\\maintest.docx'
    #Act
    parameterisedMain("https://hack.me/", testFilePath)
    if os.path.isfile(testFilePath):
        document = Document(testFilePath)

    #Assert
    assert os.path.isfile(testFilePath) == True
    #Assert headers table is populated
    assert document.tables[1].rows[1].cells[0].text == "X-Content-Type-Options"
    #Assert cookies table is populated
    assert document.tables[3].rows[1].cells[0].text == "Samesite"

    #Teardown
    os.remove(testFilePath)